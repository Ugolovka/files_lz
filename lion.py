import re
from collections import Counter
import pandas as pd
from docx import Document
import matplotlib.pyplot as plt

#чтение текста из .docx файла
def read_docx(file_path):
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

#анализ текста
def analyze_text_from_docx(file_path):
    text = read_docx(file_path).lower()
    
    #удаление лишних символов и разделение на слова
    words = re.findall(r'\b\w+\b', text)
    
    #подсчет частоты слов через контер
    word_count = Counter(words)
    total_words = sum(word_count.values())
    
    #подсчет частоты букв
    letters = re.findall(r'[а-яёa-z]', text)  #может в тексте есть другие буквы кроме русских поэтому ещё и подсчёт английских
    letter_count = Counter(letters)
    total_letters = sum(letter_count.values())
    
    #форматирование таблицы по словам
    word_stats = [
        {"Слово": word, "Частота встреч в раз": count, "Частота встреч в %": round((count / total_words) * 100, 2)}
        for word, count in word_count.items()
    ]
    word_df = pd.DataFrame(word_stats)
    
    #форматирование таблицы по буквам
    letter_stats = [
        {"Буква": letter, "Частота встреч в раз": count, "Частота встреч в %": round((count / total_letters) * 100, 2)}
        for letter, count in letter_count.items()
    ]
    letter_df = pd.DataFrame(letter_stats)
    
    return word_df, letter_df

#построение гистограммы частоты букв
def plot_letter_histogram(letter_df):
    plt.figure(figsize=(10, 6))
    plt.bar(letter_df["Буква"], letter_df["Частота встреч в %"], color="skyblue")
    plt.xlabel("Буквы", fontsize=12)
    plt.ylabel("Частота встречаемости (%)", fontsize=12)
    plt.title("Частота встречаемости букв", fontsize=14)
    plt.grid(axis="y", linestyle="--", alpha=0.7)
    plt.show()

#построение гистограммы частоты слов
def plot_word_histogram(word_df, top_n=10):
    #сортировка по частоте и выбор top_n слов
    top_words = word_df.sort_values(by="Частота встреч в раз", ascending=False).head(top_n)
    
    plt.figure(figsize=(10, 6))
    plt.bar(top_words["Слово"], top_words["Частота встреч в раз"], color="lightgreen")
    plt.xlabel("Слова")
    plt.ylabel("Частота встречаемости (раз)")
    plt.title(f"Топ-{top_n} наиболее частых слов")
    plt.xticks(rotation=45)
    plt.grid(axis="y", linestyle="--")
    plt.show()

file_path = "lion.docx" 
word_df, letter_df = analyze_text_from_docx(file_path)

#построение гистограмм
print("Построение гистограммы для букв...")
plot_letter_histogram(letter_df)

print("Построение гистограммы для слов...")
plot_word_histogram(word_df, top_n=10)



"""from docx import Document   варик без функций 
import matplotlib.pyplot as plt
from collections import Counter

file_path = "example.docx"

#читаем документ
doc = Document(file_path)

#извлекаем текст из всех абзацев
text = " ".join([paragraph.text for paragraph in doc.paragraphs])

#подсчёт слов
words = text.split()
word_count = len(words)

#подсчёт букв
letters = [char for char in text if char.isalpha()]
letter_count = len(letters)

#статистика слов
word_frequency = Counter(words)

#статистика букв
letter_frequency = Counter(letters)

#построение гистограммы по словам
plt.figure(figsize=(10, 5))
plt.bar(word_frequency.keys(), word_frequency.values(), color='skyblue')
plt.title('Частота слов')
plt.xlabel('Слова')
plt.ylabel('Частота')
plt.xticks(rotation=45, ha='right', fontsize=8)
plt.tight_layout()
plt.show()

#построение гистограммы по буквам
plt.figure(figsize=(10, 5))
plt.bar(letter_frequency.keys(), letter_frequency.values(), color='lightgreen')
plt.title('Частота букв')
plt.xlabel('Буквы')
plt.ylabel('Частота')
plt.xticks(rotation=0)
plt.tight_layout()
plt.show()

#вывод
print(f"Количество слов: {word_count}")
print(f"Количество букв: {letter_count}")"""
